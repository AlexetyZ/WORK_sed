import docx
import re

doc = docx.Document("C:\\Users\\user\\PycharmProjects\\django3\\SPD\\SPD\\apps\\proverka\\doc_templates\\шаблон протокола об административных правонарушениях.docx")


# print(doc.paragraphs[9].text)
RegEx = {}
for par in doc.paragraphs:
    if re.findall(r'{\s(\S+)\s}', par.text):
        # print(re.findall(r'\{\s(\S+)\s\}', par.text))
        reg_in_par = re.findall(r'{\s(\S+)\s}', par.text)
        for reg in reg_in_par:
            try:
                reg = reg.split('.')[0]
                RegEx[reg] = reg
            except Exception as ex:
                # print(ex)
                RegEx[reg] = reg
try:
    tab = doc.tables[0].rows[0].cells[0].tables[0].rows[0].cells[0].text
except:
    try:
        tab = doc.tables[0].rows[0].cells[0].text
    except:
        tab = doc.paragraphs

for tab in doc.tables:

    for row in tab.rows:
        for cell in row.cells:
            for par in cell.paragraphs:

                if re.findall(r'{\s(\S+)\s}', par.text):
                    # print(re.findall(r'\{\s(\S+)\s\}', par.text))
                    reg_in_par = re.findall(r'{\s(\S+)\s}', par.text)
                    for reg in reg_in_par:
                        try:
                            reg = reg.split('.')[0]
                        except Exception as ex:
                            # print(ex)
                            RegEx[reg] = reg
            for celltables in cell.tables:
                for parrow in celltables.rows:
                    for parcell in parrow.cells:
                        if re.findall(r'{\s(\S+)\s}', parcell.text):
                            # print(re.findall(r'\{\s(\S+)\s\}', par.text))
                            reg_in_par = re.findall(r'{\s(\S+)\s}', parcell.text)
                            for reg in reg_in_par:
                                try:
                                    reg = reg.split('.')[0]
                                except Exception as ex:
                                    # print(ex)
                                    RegEx[reg] = reg

print(RegEx)
print(len(RegEx))
# for key, value in RegEx.items():
#
#   print("'{0}': {1},".format(key,value))
