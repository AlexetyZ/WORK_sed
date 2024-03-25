import docx


doc = docx.Document()


def addPic(cell, picPath, width: int = 1800720, height: int = 1800720):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(picPath, width=width, height=height)


contentKind = {
    '<image>': {'func': addPic}
}


def createTable(*datas):
    lenght = len(datas[0])
    mainTable = doc.add_table(cols=lenght, rows=1)
    for data in datas:
        lastRow = mainTable.add_row().cells
        # print(data)
        for cell, d in zip(lastRow, data):
            for ck, vals in contentKind.items():
                if ck in d:
                    picName = str(d).split(ck)[1]
                    vals['func'](cell, picName)
                else:
                    cell.text = d


def save_doc(name):
    doc.save(name)


def main():
    datas = [['<image>downloadedFiles/c28de5bd-0339-4ac2-8863-b54491ed16a3.jpg', 'Башанкаев Бадма Николаевич', 'Председатель комитета Государственной Думы по охране здоровья']]
    createTable(*datas)
    save_doc('example.docx')


if __name__ == '__main__':
    main()
